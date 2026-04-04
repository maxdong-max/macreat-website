#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成中文版 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 确保输出目录存在
os.makedirs('/Users/openclaw/Downloads', exist_ok=True)

def add_heading(doc, text, level=0):
    """添加标题"""
    if level == 0:
        heading = doc.add_heading(text, level=1)
    else:
        heading = doc.add_heading(text, level=level + 1)
    return heading

def add_paragraph(doc, text, bold=False):
    """添加段落"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_code_block(doc, code):
    """添加代码块（保持英文）"""
    p = doc.add_paragraph()
    p.add_run(code).font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

# ============ 文档1: 项目说明 ============
doc1 = Document()

# 标题
title = doc1.add_heading('Macreat 网站升级项目说明', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc1, '\n项目结构\n', bold=True)

# 项目结构树
structure = """
.
├── macreat-frontend/          # Next.js 前端 (Vercel)
│   ├── components/
│   │   ├── Header.js          # 导航栏
│   │   ├── Footer.js         # 页脚
│   │   └── FloatingSocial.js # WhatsApp/询盘悬浮按钮
│   ├── pages/
│   │   ├── _app.js           # 应用包装器（含分析）
│   │   └── index.js          # 首页
│   ├── lib/
│   │   ├── content.js        # 100%保留的网站内容
│   │   ├── seo.js            # SEO 配置
│   │   └── analytics.js      # 分析跟踪
│   ├── styles/
│   │   └── globals.css       # 响应式 CSS 样式
│   ├── next.config.js        # Next.js 配置
│   └── package.json
│
├── macreat-backend/           # Express.js 后端 (Railway)
│   ├── models/
│   │   ├── Analytics.js      # 访客跟踪 schema
│   │   ├── Content.js        # 动态内容 schema
│   │   ├── Chat.js           # AI 聊天会话 schema
│   │   └── User.js           # 管理员用户 schema
│   ├── routes/
│   │   ├── analytics.js      # 分析 API
│   │   ├── content.js        # 内容管理 API
│   │   ├── auth.js           # 认证 API
│   │   ├── chat.js           # AI 聊天 API
│   │   └── seo.js            # SEO 配置 API
│   ├── services/
│   │   └── chatService.js    # AI 回复逻辑
│   ├── middleware/
│   │   └── auth.js           # JWT 认证
│   ├── server.js             # Express 服务器
│   ├── package.json
│   └── .env.example          # 环境变量模板
│
├── DEPLOYMENT.md             # 部署指南
└── README.md                 # 本文件
"""
p = doc1.add_paragraph()
p.add_run(structure).font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

add_heading(doc1, '建设内容', level=1)

add_paragraph(doc1, '\n1. 前端 (Next.js)\n', bold=True)
add_paragraph(doc1, '✅ 100% 保留所有原始内容')
add_paragraph(doc1, '✅ 响应式设计（移动端/平板/桌面端）')
add_paragraph(doc1, '✅ 多语言支持（英语/中文/印尼语）')
add_paragraph(doc1, '✅ SEO 优化（含 meta 标签）')
add_paragraph(doc1, '✅ 分析跟踪集成')
add_paragraph(doc1, '✅ 悬浮 WhatsApp/询盘按钮')
add_paragraph(doc1, '✅ 联系表单及 API 端点')

add_paragraph(doc1, '\n2. 后端 (Express.js)\n', bold=True)
add_paragraph(doc1, '✅ 用户认证（JWT）')
add_paragraph(doc1, '✅ 内容管理（编辑任意文本/图片）')
add_paragraph(doc1, '✅ 分析仪表盘数据')
add_paragraph(doc1, '✅ AI 客服（多语言）')
add_paragraph(doc1, '✅ SEO 配置 API')
add_paragraph(doc1, '✅ MongoDB 数据库集成')

add_paragraph(doc1, '\n3. 功能特点\n', bold=True)
add_paragraph(doc1, '• 多语言：支持英语/中文/印尼语内容编辑')
add_paragraph(doc1, '• AI 客服：自动回复产品/价格咨询')
add_paragraph(doc1, '• 分析：实时访客跟踪')
add_paragraph(doc1, '• SEO：自动生成站点地图、地理定位')
add_paragraph(doc1, '• 社交链接：WhatsApp、Facebook、Telegram 等')

add_heading(doc1, '内容保留', level=1)
add_paragraph(doc1, '所有原始网站文本均保留在 lib/content.js 中：')
add_paragraph(doc1, '• 英雄区域文本')
add_paragraph(doc1, '• 产品名称和规格')
add_paragraph(doc1, '• 案例研究描述')
add_paragraph(doc1, '• 新闻文章')
add_paragraph(doc1, '• 常见问题')
add_paragraph(doc1, '• 联系方式')
add_paragraph(doc1, '• 页脚链接')
add_paragraph(doc1, '\n未修改或删除任何内容。', bold=True)

add_heading(doc1, '部署说明', level=1)
add_paragraph(doc1, '此代码需要部署后才能使用：')
add_paragraph(doc1, '1. 前端 → Vercel')
add_paragraph(doc1, '2. 后端 → Railway')
add_paragraph(doc1, '3. 数据库 → MongoDB（Railway 或 Atlas）')
add_paragraph(doc1, '\n详细部署说明请参阅 DEPLOYMENT.md')

# 保存文档1
doc1.save('/Users/openclaw/Downloads/Macreat_网站升级项目说明.docx')
print('✅ 已生成: Macreat_网站升级项目说明.docx')


# ============ 文档2: 部署教程 ============
doc2 = Document()

# 标题
title2 = doc2.add_heading('Macreat 网站部署教程', 0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc2, '\n概述\n', bold=True)
add_paragraph(doc2, '本指南涵盖升级后的 macreat.com 网站部署，包括：')
add_paragraph(doc2, '• 前端：Next.js (Vercel)')
add_paragraph(doc2, '• 后端：Express.js (Railway)')
add_paragraph(doc2, '• 数据库：MongoDB (Railway)')
add_paragraph(doc2, '• 功能：多语言、AI 客服、分析、SEO')

add_heading(doc2, '准备工作', level=1)
add_paragraph(doc2, '1. GitHub/GitLab 账户 - 用于版本控制')
add_paragraph(doc2, '2. Vercel 账户 - 用于前端托管')
add_paragraph(doc2, '3. Railway 账户 - 用于后端托管')
add_paragraph(doc2, '4. MongoDB Atlas 或 Railway MongoDB - 用于数据库')

add_heading(doc2, '步骤一：准备前端', level=1)

add_paragraph(doc2, '\n1.1 更新配置', bold=True)
add_paragraph(doc2, '编辑 next.config.js，更新 API 重写 URL：')

code1 = """async rewrites() {
  return [
    {
      source: '/api/:path*',
      destination: 'https://your-railway-app.up.railway.app/api/:path*',
    },
  ];
}"""
add_code_block(doc2, code1)
add_paragraph(doc2, '将 your-railway-app 替换为您的 Railway 项目名称。')

add_paragraph(doc2, '\n1.2 构建并部署到 Vercel', bold=True)

code2 = """cd macreat-frontend

# 安装依赖
npm install

# 本地测试构建
npm run build

# 部署到 Vercel
# 选项一：通过 CLI
npm i -g vercel
vercel login
vercel --prod

# 选项二：通过 GitHub
# 推送到 GitHub，从仪表盘连接到 Vercel"""
add_code_block(doc2, code2)

add_paragraph(doc2, '\n1.3 在 Vercel 中配置环境变量', bold=True)
add_paragraph(doc2, '进入 Vercel 仪表盘 → 设置 → 环境变量：')

code3 = """NEXT_PUBLIC_API_URL=https://your-railway-app.up.railway.app
NEXT_PUBLIC_SITE_URL=https://macreat.com"""
add_code_block(doc2, code3)

add_heading(doc2, '步骤二：准备后端', level=1)

add_paragraph(doc2, '\n2.1 更新环境变量', bold=True)

code4 = """cd macreat-backend
cp .env.example .env"""
add_code_block(doc2, code4)
add_paragraph(doc2, '编辑 .env：')

code5 = """PORT=3000
MONGODB_URI=your-mongodb-connection-string
JWT_SECRET=your-secure-random-string
FRONTEND_URL=https://macreat.com"""
add_code_block(doc2, code5)

add_paragraph(doc2, '\n2.2 部署到 Railway', bold=True)

code6 = """# 安装 Railway CLI
npm i -g @railway/cli

# 登录
railway login

# 初始化项目
railway init

# 添加 MongoDB 插件
railway add mongodbatlas

# 部署"""
add_code_block(doc2, code6)

add_paragraph(doc2, '\n2.3 获取 Railway URL', bold=True)
add_paragraph(doc2, '部署后，获取您的 Railway 应用 URL：')

code7 = """railway domain"""
add_code_block(doc2, code7)

add_heading(doc2, '步骤三：连接前端到后端', level=1)

add_paragraph(doc2, '\n3.1 更新 API URL', bold=True)
add_paragraph(doc2, '在 Vercel 中更新环境变量：')

code8 = """NEXT_PUBLIC_API_URL=https://your-railway-app.up.railway.app"""
add_code_block(doc2, code8)

add_paragraph(doc2, '\n3.2 测试 API 连接', bold=True)
add_paragraph(doc2, '访问：https://your-railway-app.up.railway.app/api/health')
add_paragraph(doc2, '应返回：{"status":"ok","timestamp":"..."}')

add_heading(doc2, '步骤四：域名配置', level=1)

add_paragraph(doc2, '\n4.1 连接自定义域名到 Vercel', bold=True)
add_paragraph(doc2, '1. 进入 Vercel 仪表盘 → 设置 → 域名')
add_paragraph(doc2, '2. 添加 macreat.com')
add_paragraph(doc2, '3. 按照 DNS 配置说明操作')

add_paragraph(doc2, '\n4.2 DNS 设置', bold=True)
add_paragraph(doc2, '在您的域名服务商处添加以下记录：')

# 表格
table = doc2.add_table(rows=3, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '类型'
hdr_cells[1].text = '名称'
hdr_cells[2].text = '值'
row1 = table.rows[1].cells
row1[0].text = 'CNAME'
row1[1].text = 'www'
row1[2].text = 'cdn.vercel.com'
row2 = table.rows[2].cells
row2[0].text = 'A'
row2[1].text = '@'
row2[2].text = '76.76.21.21'

add_paragraph(doc2, '\n4.3 SSL', bold=True)
add_paragraph(doc2, 'Vercel 自动配置 SSL。请确保在设置中强制启用 HTTPS。')

add_heading(doc2, '步骤五：管理面板设置', level=1)

add_paragraph(doc2, '\n5.1 创建管理员账户', bold=True)
add_paragraph(doc2, '调用注册端点一次：')

code9 = """curl -X POST https://your-railway-app.up.railway.app/api/auth/register \\
  -H "Content-Type: application/json" \\
  -d '{
    "username": "admin",
    "password": "your-secure-password",
    "email": "admin@macreat.com"
  }'"""
add_code_block(doc2, code9)

add_paragraph(doc2, '\n5.2 访问管理面板', bold=True)
add_paragraph(doc2, '管理面板地址：https://macreat.com/admin')

add_heading(doc2, '步骤六：功能验证', level=1)

add_paragraph(doc2, '\n6.1 多语言', bold=True)
add_paragraph(doc2, '• 访问 /、/zh、/id')
add_paragraph(doc2, '• 检查导航栏中的语言切换器')

add_paragraph(doc2, '\n6.2 分析', bold=True)
add_paragraph(doc2, '• 访问 https://your-railway-app.up.railway.app/api/analytics/summary')
add_paragraph(doc2, '• 应显示页面浏览量和访客数据')

add_paragraph(doc2, '\n6.3 AI 聊天', bold=True)
add_paragraph(doc2, '• 在前端测试聊天小部件')
add_paragraph(doc2, '• 在管理面板中查看聊天历史')

add_paragraph(doc2, '\n6.4 内容管理', bold=True)
add_paragraph(doc2, '• 登录管理面板')
add_paragraph(doc2, '• 尝试编辑某个板块')
add_paragraph(doc2, '• 验证更改出现在前端')

add_paragraph(doc2, '\n6.5 SEO', bold=True)
add_paragraph(doc2, '• 检查站点地图：https://macreat.com/api/seo/sitemap.xml')
add_paragraph(doc2, '• 检查 robots.txt：https://macreat.com/api/seo/robots.txt')
add_paragraph(doc2, '• 验证页面源代码中的 meta 标签')

add_heading(doc2, '问题排查', level=1)

add_paragraph(doc2, '\n前端问题', bold=True)
add_paragraph(doc2, 'API 无法连接')
add_paragraph(doc2, '• 验证 Vercel 中的 NEXT_PUBLIC_API_URL')
add_paragraph(doc2, '• 检查 Railway 是否在运行（未停止）')
add_paragraph(doc2, '• 直接测试健康端点')
add_paragraph(doc2, '\n构建错误')
add_paragraph(doc2, '• 清除 Vercel 缓存并重新构建')
add_paragraph(doc2, '• 检查 Node.js 版本兼容性')

add_paragraph(doc2, '\n后端问题', bold=True)
add_paragraph(doc2, 'MongoDB 连接失败')
add_paragraph(doc2, '• 验证 .env 中的连接字符串')
add_paragraph(doc2, '• 检查 MongoDB Atlas 中的 IP 白名单')
add_paragraph(doc2, '\n认证错误')
add_paragraph(doc2, '• 确保 JWT_SECRET 已设置')
add_paragraph(doc2, '• 清除浏览器 cookies 并重新登录')

add_heading(doc2, '维护', level=1)

add_paragraph(doc2, '\n数据库备份', bold=True)

code10 = """# 从 MongoDB Atlas 导出数据
mongodump --uri="your-mongodb-uri" --out=./backup"""
add_code_block(doc2, code10)

add_paragraph(doc2, '\n更新前端', bold=True)
add_paragraph(doc2, '• 进行更改，提交并推送到 GitHub')
add_paragraph(doc2, '• Vercel 会在推送到 main 分支时自动部署')

add_paragraph(doc2, '\n更新后端', bold=True)
add_paragraph(doc2, '• 进行更改')
add_paragraph(doc2, '• railway up')
add_paragraph(doc2, '• Railway 会自动部署')

add_heading(doc2, '支持', level=1)
add_paragraph(doc2, '如遇问题，请检查：')
add_paragraph(doc2, '1. Railway 日志：railway logs')
add_paragraph(doc2, '2. Vercel 函数日志')
add_paragraph(doc2, '3. 浏览器控制台中的前端错误')

add_heading(doc2, '费用估算（月度）', level=1)

# 表格
table2 = doc2.add_table(rows=4, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = '服务'
hdr_cells2[1].text = '免费套餐'
hdr_cells2[2].text = '付费套餐'
row1 = table2.rows[1].cells
row1[0].text = 'Vercel'
row1[1].text = '100GB 流量'
row1[2].text = '约 $20/100GB'
row2 = table2.rows[2].cells
row2[0].text = 'Railway'
row2[1].text = '$5 积分/月'
row2[2].text = '约 $10-20'
row3 = table2.rows[3].cells
row3[0].text = 'MongoDB Atlas'
row3[1].text = '512MB'
row3[2].text = '约 $10-25'

add_paragraph(doc2, '\n总计：生产环境使用约 $25-65/月')

# 保存文档2
doc2.save('/Users/openclaw/Downloads/Macreat_网站部署教程.docx')
print('✅ 已生成: Macreat_网站部署教程.docx')

print('\n🎉 所有文档生成完成！')